from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from accounts.decorators import admin_required
from .models import *
from django.core.paginator import Paginator
from django.db.models import Q
from django.http import JsonResponse, HttpResponse
from django.template.loader import render_to_string, get_template
import csv
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from io import BytesIO
from xhtml2pdf import pisa
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .forms import *


@login_required
@admin_required
def add_employee(request, company_name):
    company = get_object_or_404(Company, name__iexact=company_name)

    if request.method == "POST":
        form = EmployeeForm(request.POST, request.FILES, company=company)
        if form.is_valid():
            employee = form.save(commit=False)
            employee.company = company
            employee.created_by = request.user
            employee.save()
            return redirect('dashboard')
    else:
        form = EmployeeForm(company=company)

    return render(request, 'employees/add_employee.html', {
        'form': form,
        'company': company
    })


# @login_required
# def employee_list(request):
#     employees = Employee.objects.select_related('team', 'company')

#     status = request.GET.get('status')
#     if status in ['active', 'resigned']:
#         employees = employees.filter(status=status)

#     return render(request, 'employees/employee_list.html', {
#         'employees': employees
#     })

@login_required
def company_list(request):
    companies = Company.objects.all()
    return render(request, 'employees/company_list.html', {'companies': companies})

@login_required
def company_employees(request, company_id):
    company = get_object_or_404(Company, id=company_id)
    employees = Employee.objects.filter(company=company).select_related('team')

    # Get filter parameters
    status = request.GET.get('status', '')
    team = request.GET.get('team', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    search = request.GET.get('search', '')

    # Apply filters
    if status:
        employees = employees.filter(status=status)

    if team:
        employees = employees.filter(team_id=team)

    if date_from:
        employees = employees.filter(date_of_joining__gte=date_from)
    if date_to:
        employees = employees.filter(date_of_joining__lte=date_to)

    if search:
        employees = employees.filter(
            Q(name__icontains=search) |
            Q(email__icontains=search) |
            Q(role__icontains=search)
        )
        # 🔹 EXPORT HANDLING (BEFORE PAGINATION!)
    export_type = request.GET.get('export')
    if export_type:
        return handle_export(request, employees, company, export_type)

    # 🔹 PAGINATION
    paginator = Paginator(employees, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Handle AJAX request for live filtering
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        html = render_to_string('employees/partials/employee_rows.html', {
            'employees': employees,
            'request': request
        })
        pagination_html = render_to_string('employees/partials/pagination.html', {
            'page_obj': page_obj
        })
        return JsonResponse({
            'html': html,
            'pagination': pagination_html,
            'count': paginator.count
        })

    # Handle exports
    

    teams = Team.objects.filter(company=company)

    return render(request, 'employees/company_employees.html', {
        'company': company,
        'employees': employees,
        'teams': teams,
        'page_obj': page_obj
    })


def handle_export(request, employees, company, export_type):
    """Handle all export types"""
    
    if export_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{company.name}_employees.csv"'
        writer = csv.writer(response)
        writer.writerow(['Name', 'Team', 'Role', 'Date Of Joining', 'Status'])

        for emp in employees:
            writer.writerow([
                emp.name,
                emp.team.name if emp.team else 'N/A',
                emp.role,
                emp.date_of_joining.strftime('%Y-%m-%d'),
                emp.status.title()
            ])
        return response

    elif export_type == 'excel':
        wb = Workbook()
        ws = wb.active
        ws.title = "Employees"

        # Header styling
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)
        
        # Add headers
        headers = ['Name', 'Team', 'Role', 'Date Of Joining', 'Status']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Add data
        for emp in employees:
            ws.append([
                emp.name,
                emp.team.name if emp.team else 'N/A',
                emp.role,
                emp.date_of_joining.strftime('%Y-%m-%d'),
                emp.status.title()
            ])

        # Adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{company.name}_employees.xlsx"'
        wb.save(response)
        return response

    elif export_type == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{company.name}_employees.pdf"'
        
        doc = SimpleDocTemplate(response, pagesize=A4)
        elements = []
        
        styles = getSampleStyleSheet()
        title = Paragraph(f"{company.name} - Employee List", styles['Heading1'])
        elements.append(title)
        
        # Table data
        data = [['Name', 'Team', 'Role', 'Date Of Joining', 'Status']]
        
        for emp in employees:
            data.append([
                emp.name,
                emp.team.name if emp.team else 'N/A',
                emp.role,
                emp.date_of_joining.strftime('%Y-%m-%d'),
                emp.status.title()
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
        ]))
        
        elements.append(table)
        doc.build(elements)
        return response

    elif export_type == 'word':
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{company.name}_employees.docx"'
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{company.name} - Employee List', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line break
        doc.add_paragraph()
        
        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Name', 'Team', 'Role', 'Date Of Joining', 'Status']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Data rows
        for emp in employees:
            row_cells = table.add_row().cells
            row_cells[0].text = emp.name
            row_cells[1].text = emp.team.name if emp.team else 'N/A'
            row_cells[2].text = emp.role
            row_cells[3].text = emp.date_of_joining.strftime('%Y-%m-%d')
            row_cells[4].text = emp.status.title()
        
        doc.save(response)
        return response

    elif export_type == 'print':
        return render(request, 'employees/exports/employees_print.html', {
            'employees': employees,
            'company': company
        })



@login_required
@admin_required
def edit_employee(request, employee_id):
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        form = EmployeeForm(
            request.POST,
            request.FILES,
            instance=employee,
            company=employee.company
        )
        if form.is_valid():
            form.save()
            return redirect('company_employees', employee.company.id)
    else:
        form = EmployeeForm(instance=employee)
       
        form.fields['team'].queryset = Team.objects.filter(
            company=employee.company
        )

    return render(request, 'employees/edit_employee.html', {
        'form': form,
        'employee': employee
    })

@login_required
@admin_required
def delete_employee(request, employee_id):
    employee = get_object_or_404(Employee, id=employee_id)

    if request.method == 'POST':
        employee.status = 'resigned'
        employee.save()
        return redirect('company_list')

    return render(request, 'employees/delete_employee.html', {
        'employee': employee
    })

@login_required
@admin_required
def add_company(request):
    if request.method == 'POST':
        form = CompanyForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('dashboard')
    else:
        form = CompanyForm()

    return render(request, 'employees/add_company.html', {'form': form})

@login_required
@admin_required
def view_company(request):
    company = Company.objects.all()
    return render(request,"employees/view_company.html",{'company':company})

@login_required
@admin_required
def edit_company(request,pk):
    company=get_object_or_404(Company,pk=pk)
    if request.method == "POST":
        form = CompanyForm(request.POST,instance=company)
        if form.is_valid():
            form.save()
            return redirect('view_company')
    else:
        form = CompanyForm(instance=company)
    return render(request,"employees/edit_company.html",{'form':form})

@login_required
@admin_required
def company_delete(request,pk):
    detail=Company.objects.get(id=pk)
    detail.delete()
    return redirect('view_company')

@login_required
@admin_required
def add_team(request):
    if request.method == 'POST':
        form = TeamForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('dashboard')
    else:
        form = TeamForm()

    return render(request, 'employees/add_team.html', {'form': form})

@login_required
@admin_required
def view_team(request):
    team = Team.objects.all()
    return render(request,"employees/view_team.html",{'team':team})